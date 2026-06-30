"""
src/phase1/module4_parsers.py
==============================
MODULE 4 – Parser Selection

Responsibility
--------------
READ the file and return its raw content.  That is the only job of
every parser in this module.  Parsers NEVER:
  - Perform NLP
  - Clean or normalise data
  - Extract candidate fields
  - Merge or match candidates

Parser outputs
--------------
  ParsedContent
    .data_category  : "structured" | "unstructured"
    .content        : list[dict] (structured) or str (unstructured)
    .parser_used    : parser class name
    .raw_length     : row count (structured) or character count (unstructured)
    .parse_warning  : non-fatal notice (e.g. "pypdf not installed")

All parsers are instantiated once by ParserRegistry (singleton pattern).
"""

from __future__ import annotations

import csv
import json
import logging
import re
import zipfile
from abc import ABC, abstractmethod
from pathlib import Path

from src.core.config_loader import get_config
from src.core.exceptions import NoParserRegisteredError, ParsingFailedError
from src.core.models import FileMetadata, ParsedContent

logger = logging.getLogger("phase1.module4_parsers")


# ──────────────────────────────────────────────────────────────
# Parser base
# ──────────────────────────────────────────────────────────────

class BaseParser(ABC):
    @abstractmethod
    def parse(self, path: Path) -> ParsedContent:
        raise NotImplementedError


# ──────────────────────────────────────────────────────────────
# Structured parsers
# ──────────────────────────────────────────────────────────────

class CsvParser(BaseParser):
    """Returns a list of dicts (one per row) via csv.DictReader."""

    def parse(self, path: Path) -> ParsedContent:
        try:
            with open(path, "r", encoding="utf-8", errors="replace", newline="") as f:
                reader = csv.DictReader(f)
                rows = [dict(row) for row in reader]
        except csv.Error as exc:
            raise ParsingFailedError(f"CSV parse error in {path.name}: {exc}") from exc

        logger.info("CsvParser: %d rows from %s", len(rows), path.name)
        return ParsedContent(
            data_category="structured",
            content=rows,
            parser_used="CsvParser",
            raw_length=len(rows),
        )


class JsonParser(BaseParser):
    """
    Returns the parsed JSON object.
    If the top-level object is a dict (single candidate), it is wrapped
    in a list so downstream code always sees list[dict].
    """

    def parse(self, path: Path) -> ParsedContent:
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                data = json.load(f)
        except json.JSONDecodeError as exc:
            raise ParsingFailedError(f"JSON parse error in {path.name}: {exc}") from exc

        if isinstance(data, dict):
            data = [data]
        length = len(data) if isinstance(data, list) else 1
        logger.info("JsonParser: %d records from %s", length, path.name)
        return ParsedContent(
            data_category="structured",
            content=data,
            parser_used="JsonParser",
            raw_length=length,
        )


# ──────────────────────────────────────────────────────────────
# Unstructured parsers – all return a plain text string
# ──────────────────────────────────────────────────────────────

class TxtParser(BaseParser):
    def parse(self, path: Path) -> ParsedContent:
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            raise ParsingFailedError(f"Cannot read TXT {path.name}: {exc}") from exc

        logger.info("TxtParser: %d chars from %s", len(text), path.name)
        return ParsedContent(
            data_category="unstructured",
            content=text,
            parser_used="TxtParser",
            raw_length=len(text),
        )


class PdfParser(BaseParser):
    """
    Uses `pypdf` for text extraction.
    Falls back gracefully (with a warning) if pypdf is not installed.
    """

    def parse(self, path: Path) -> ParsedContent:
        try:
            from pypdf import PdfReader
        except ImportError:
            size = path.stat().st_size
            logger.warning("pypdf not installed — PDF text not extracted: %s", path.name)
            return ParsedContent(
                data_category="unstructured",
                content="",
                parser_used="PdfParser(no-pypdf-fallback)",
                raw_length=size,
                parse_warning=(
                    "pypdf is not installed. "
                    "Run `pip install pypdf` for real PDF text extraction."
                ),
            )

        try:
            reader = PdfReader(str(path))
            pages: list[str] = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                # pypdf sometimes inserts spurious whitespace between characters.
                # Collapse runs of spaces that exceed 3 consecutive spaces to 1.
                page_text = re.sub(r" {4,}", " ", page_text)
                pages.append(page_text)
            text = "\n".join(pages).strip()
        except Exception as exc:
            raise ParsingFailedError(f"PDF read error in {path.name}: {exc}") from exc

        logger.info("PdfParser: %d chars from %s (%d pages)", len(text), path.name, len(reader.pages))
        return ParsedContent(
            data_category="unstructured",
            content=text,
            parser_used="PdfParser",
            raw_length=len(text),
        )


class DocxParser(BaseParser):
    """
    Extracts text from a DOCX file.
    Tries `python-docx` first; falls back to raw XML extraction (stdlib only).
    Paragraph order is preserved exactly as in the document.
    """

    WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def parse(self, path: Path) -> ParsedContent:
        # Attempt python-docx (preserves more structure)
        try:
            import docx as _docx

            doc = _docx.Document(str(path))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract table cell text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in paragraphs:
                            paragraphs.append(cell_text)
            text = "\n".join(paragraphs)
            parser_name = "DocxParser(python-docx)"

        except ImportError:
            # Fallback: raw XML
            text = self._extract_via_xml(path)
            parser_name = "DocxParser(xml-fallback)"
        except Exception as exc:
            raise ParsingFailedError(f"DOCX read error in {path.name}: {exc}") from exc

        logger.info("%s: %d chars from %s", parser_name, len(text), path.name)
        return ParsedContent(
            data_category="unstructured",
            content=text,
            parser_used=parser_name,
            raw_length=len(text),
        )

    def _extract_via_xml(self, path: Path) -> str:
        import xml.etree.ElementTree as ET

        try:
            with zipfile.ZipFile(path) as z:
                xml_bytes = z.read("word/document.xml")
        except (zipfile.BadZipFile, KeyError) as exc:
            raise ParsingFailedError(
                f"Cannot read DOCX XML content from {path.name}: {exc}"
            ) from exc

        try:
            root = ET.fromstring(xml_bytes)
        except ET.ParseError as exc:
            raise ParsingFailedError(
                f"DOCX XML parse error in {path.name}: {exc}"
            ) from exc

        paragraphs: list[str] = []
        ns = self.WORD_NS
        for para in root.iter(f"{ns}p"):
            runs = [
                node.text
                for node in para.iter(f"{ns}t")
                if node.text
            ]
            if runs:
                paragraphs.append("".join(runs))

        return "\n".join(paragraphs)


class RtfParser(BaseParser):
    """
    Minimal RTF → plain text stripper (no external dependency).
    Handles the most common RTF control words found in resume files.
    """

    _CONTROL = re.compile(r"\\([a-zA-Z]+)(-?\d+)? ?")
    _HEX_ESC = re.compile(r"\\'[0-9a-fA-F]{2}")
    _LINEBREAK = re.compile(r"\\(par|line|pard)\b ?")

    def parse(self, path: Path) -> ParsedContent:
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            raise ParsingFailedError(f"Cannot read RTF {path.name}: {exc}") from exc

        text = self._strip(raw)
        logger.info("RtfParser: %d chars from %s", len(text), path.name)
        return ParsedContent(
            data_category="unstructured",
            content=text,
            parser_used="RtfParser",
            raw_length=len(text),
        )

    def _strip(self, raw: str) -> str:
        text = self._HEX_ESC.sub("", raw)
        text = self._LINEBREAK.sub("\n", text)
        text = self._CONTROL.sub(" ", text)
        text = text.replace("{", " ").replace("}", " ")
        lines = [" ".join(ln.split()) for ln in text.splitlines()]
        return "\n".join(ln for ln in lines if ln)


class HtmlParser(BaseParser):
    """
    Extracts visible text from HTML using stdlib html.parser.
    Skips <script>, <style>, <head> and <noscript> content.
    Block-level tags introduce line breaks.
    """

    _BLOCK_TAGS = {
        "p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6",
        "section", "article", "header", "footer", "table", "ul", "ol",
    }
    _SKIP_TAGS = {"script", "style", "head", "noscript", "meta", "link"}

    def parse(self, path: Path) -> ParsedContent:
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            raise ParsingFailedError(f"Cannot read HTML {path.name}: {exc}") from exc

        text = self._extract(raw)
        logger.info("HtmlParser: %d chars from %s", len(text), path.name)
        return ParsedContent(
            data_category="unstructured",
            content=text,
            parser_used="HtmlParser",
            raw_length=len(text),
        )

    def _extract(self, raw: str) -> str:
        import html.parser as _hp

        chunks: list[str] = []
        skip_depth = [0]

        block_tags = self._BLOCK_TAGS
        skip_tags = self._SKIP_TAGS

        class _Extractor(_hp.HTMLParser):
            def handle_starttag(_, tag, attrs):  # noqa: N805
                if tag in skip_tags:
                    skip_depth[0] += 1
                if tag in block_tags:
                    chunks.append("\n")

            def handle_endtag(_, tag):  # noqa: N805
                if tag in skip_tags and skip_depth[0] > 0:
                    skip_depth[0] -= 1
                if tag in block_tags:
                    chunks.append("\n")

            def handle_data(_, data):  # noqa: N805
                if skip_depth[0] == 0 and data.strip():
                    chunks.append(data.strip())

        _Extractor().feed(raw)
        joined = " ".join(chunks)
        lines = [ln.strip() for ln in joined.replace(" \n", "\n").replace("\n ", "\n").split("\n")]
        return "\n".join(ln for ln in lines if ln)


class ImageOcrParser(BaseParser):
    """
    OCR using pytesseract + Pillow.
    Gracefully degrades (warning, empty string) if either is missing.
    """

    def parse(self, path: Path) -> ParsedContent:
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            logger.warning(
                "pytesseract/Pillow not installed — image OCR skipped: %s", path.name
            )
            return ParsedContent(
                data_category="unstructured",
                content="",
                parser_used="ImageOcrParser(no-ocr-fallback)",
                raw_length=path.stat().st_size,
                parse_warning=(
                    "pytesseract / Pillow are not installed. "
                    "Run `pip install pytesseract pillow` and install the "
                    "Tesseract binary for real OCR extraction."
                ),
            )

        try:
            with Image.open(path) as img:
                text = pytesseract.image_to_string(img)
        except Exception as exc:
            raise ParsingFailedError(f"OCR failed for {path.name}: {exc}") from exc

        logger.info("ImageOcrParser: %d chars from %s", len(text), path.name)
        return ParsedContent(
            data_category="unstructured",
            content=text,
            parser_used="ImageOcrParser",
            raw_length=len(text),
        )


# ──────────────────────────────────────────────────────────────
# Parser Registry (built from config)
# ──────────────────────────────────────────────────────────────

_PARSER_CLASSES: dict[str, type[BaseParser]] = {
    "CsvParser":       CsvParser,
    "JsonParser":      JsonParser,
    "TxtParser":       TxtParser,
    "PdfParser":       PdfParser,
    "DocxParser":      DocxParser,
    "RtfParser":       RtfParser,
    "HtmlParser":      HtmlParser,
    "ImageOcrParser":  ImageOcrParser,
}


class ParserRegistry:
    """
    Instantiates each unique parser once and maps every supported
    file-type extension to the correct instance.
    Mapping is driven by config → parsers.mapping.
    """

    def __init__(self) -> None:
        cfg = get_config()
        parser_map: dict[str, str] = cfg.parsers["mapping"]  # ext → class name

        # Build class → instance map (each class instantiated once)
        instances: dict[str, BaseParser] = {}
        self._ext_to_parser: dict[str, BaseParser] = {}

        for ext, class_name in parser_map.items():
            if class_name not in instances:
                cls = _PARSER_CLASSES.get(class_name)
                if cls is None:
                    logger.error("Unknown parser class in config: %s", class_name)
                    continue
                instances[class_name] = cls()
            self._ext_to_parser[ext] = instances[class_name]

    def get(self, file_type: str) -> BaseParser:
        parser = self._ext_to_parser.get(file_type)
        if parser is None:
            raise NoParserRegisteredError(
                f"No parser registered for file type '.{file_type}'"
            )
        return parser


# ──────────────────────────────────────────────────────────────
# ParserSelectionPipeline  (Module 4 orchestrator)
# ──────────────────────────────────────────────────────────────

class ParserSelectionPipeline:
    """
    Looks up the right parser from the registry and parses the file.
    Requires a valid, classified FileMetadata object.
    """

    def __init__(self) -> None:
        self._registry = ParserRegistry()

    def run(self, file_metadata: FileMetadata) -> ParsedContent:
        if not file_metadata.validation_status or not file_metadata.detected_file_type:
            raise ParsingFailedError(
                f"Cannot parse an invalid file: {file_metadata.file_name} "
                f"({file_metadata.validation_message})"
            )

        parser = self._registry.get(file_metadata.detected_file_type)
        logger.info(
            "Parsing %s with %s", file_metadata.file_name, type(parser).__name__
        )
        result = parser.parse(Path(file_metadata.file_path))

        if result.parse_warning:
            logger.warning(
                "Parser warning for %s: %s", file_metadata.file_name, result.parse_warning
            )

        logger.info(
            "Parsed %s → %d units via %s",
            file_metadata.file_name,
            result.raw_length,
            result.parser_used,
        )
        return result
